from io import BytesIO

from docx import Document
from src.infrastructure.rendering.docxtpl_render_engine import DocxtplRenderEngine


def _build_template_bytes() -> bytes:
    # docxtpl's {%tr %} row tags collapse their *entire* row to the bare
    # jinja tag, so the loop needs its own marker rows around a plain
    # content row rather than living inside the content row itself.
    document = Document()
    document.add_paragraph("Hola {{ nombre }}, tu contrato vence el {{ fecha_fin }}.")
    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Cantidad"
    table.cell(1, 0).text = "{%tr for item in items %}"
    table.cell(2, 0).text = "{{ item.nombre }}"
    table.cell(2, 1).text = "{{ item.cantidad }}"
    table.cell(3, 0).text = "{%tr endfor %}"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_detect_placeholders_finds_simple_variables() -> None:
    placeholders = DocxtplRenderEngine().detect_placeholders(_build_template_bytes())

    assert {"nombre", "fecha_fin", "items"} <= placeholders


def test_render_replaces_simple_placeholders() -> None:
    rendered = DocxtplRenderEngine().render(
        _build_template_bytes(),
        {"nombre": "Maria", "fecha_fin": "2027-01-01", "items": []},
    )

    text = "\n".join(p.text for p in Document(BytesIO(rendered)).paragraphs)

    assert "Hola Maria" in text
    assert "2027-01-01" in text


def test_render_expands_dynamic_table_rows() -> None:
    rendered = DocxtplRenderEngine().render(
        _build_template_bytes(),
        {
            "nombre": "Maria",
            "fecha_fin": "2027-01-01",
            "items": [
                {"nombre": "Sillas", "cantidad": "4"},
                {"nombre": "Mesas", "cantidad": "2"},
            ],
        },
    )

    table = Document(BytesIO(rendered)).tables[0]
    rows_text = [tuple(cell.text for cell in row.cells) for row in table.rows]

    assert ("Sillas", "4") in rows_text
    assert ("Mesas", "2") in rows_text
